#!/usr/bin/env python3
"""Build the Nutrition social media Word kit."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(15, 23, 42)
TEAL = RGBColor(13, 148, 136)
MUTED = RGBColor(71, 85, 105)
WHITE = RGBColor(255, 255, 255)
RULE = "E2E8F0"

OUT = Path(
    "/Users/charbelboufaddoul/Desktop/CTN-SOLUTION/INTERNAL-PROJECTS/Dietitian-System/Development/Dietitian/docs/marketing/Nutrition-social-media-kit.docx"
)


def set_run(run, *, size=11, bold=False, color=NAVY, italic=False, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, hex_color):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val.get("val", "single"))
        el.set(qn("w:sz"), val.get("sz", "4"))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), val.get("color", "CBD5E1"))
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def para(cell_or_doc, text, *, size=11, bold=False, color=NAVY, italic=False, align=None, space_after=8, space_before=0):
    p = cell_or_doc.add_paragraph() if hasattr(cell_or_doc, "add_paragraph") else cell_or_doc.paragraphs[0]
    if cell_or_doc.__class__.__name__ == "Cell" and cell_or_doc.paragraphs and not cell_or_doc.paragraphs[0].text:
        p = cell_or_doc.paragraphs[0]
        p.clear()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run(run, size=18, bold=True, color=NAVY)
    elif level == 2:
        set_run(run, size=14, bold=True, color=TEAL)
    else:
        set_run(run, size=12, bold=True, color=NAVY)
    return p


def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_run(run, size=9, bold=True, color=TEAL)
    run.font.all_caps = True
    return p


def post_box(doc, title, body, *, meta=None, visual=None, hashtags=None, cta=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "F8FAFC")
    set_cell_border(
        cell,
        top={"color": "0D9488", "sz": "12"},
        left={"color": "E2E8F0", "sz": "4"},
        right={"color": "E2E8F0", "sz": "4"},
        bottom={"color": "E2E8F0", "sz": "4"},
    )

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(title)
    set_run(r, size=12, bold=True, color=NAVY)

    if meta:
        mp = cell.add_paragraph()
        mp.paragraph_format.space_after = Pt(8)
        mr = mp.add_run(meta)
        set_run(mr, size=9, italic=True, color=MUTED)

    for block in body.split("\n\n"):
        bp = cell.add_paragraph()
        bp.paragraph_format.space_after = Pt(8)
        bp.paragraph_format.space_before = Pt(0)
        br = bp.add_run(block.strip("\n"))
        set_run(br, size=11, color=NAVY)

    extras = []
    if visual:
        extras.append(("Visual", visual))
    if cta:
        extras.append(("CTA", cta))
    if hashtags:
        extras.append(("Hashtags", hashtags))
    for label, value in extras:
        ep = cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(2)
        er = ep.add_run(f"{label}: ")
        set_run(er, size=10, bold=True, color=TEAL)
        vr = ep.add_run(value)
        set_run(vr, size=10, color=MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, size=11, color=NAVY)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.clear()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run(run, size=11, color=NAVY)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "0F172A")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(h)
        set_run(r, size=10, bold=True, color=WHITE)
    for r_i, row in enumerate(rows):
        fill = "F8FAFC" if r_i % 2 == 0 else "FFFFFF"
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            shade(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(str(val))
            set_run(run, size=10, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = NAVY

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run("NUTRITION")
    set_run(r, size=12, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Social media kit")
    set_run(r, size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LinkedIn · X · Instagram · Facebook · 30-day calendar")
    set_run(r, size=12, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(
        "Copy each grey box as the post. Replace [SITE] with the live URL. "
        "Do not attach a real dashboard screenshot. Do not claim a payment gateway."
    )
    set_run(r, size=11, italic=True, color=MUTED)

    add_heading_custom(doc, "How to use this file", 1)
    numbered(doc, "Pick one box. Copy only the post text, not the Visual / CTA / Hashtags lines unless you want them in the caption.")
    numbered(doc, "Paste into the platform. Add the visual described under Visual.")
    numbered(doc, "Put hashtags in the first comment on LinkedIn and Instagram when you want a cleaner caption.")
    numbered(doc, "One idea per post. One CTA: start the trial, or see how it works.")
    numbered(doc, "Audience is dietitians first. Patients second. Avoid “lose weight fast” language.")

    add_heading_custom(doc, "Positioning (keep this consistent)", 1)
    post_box(
        doc,
        "One-liner",
        "A clinic workspace for dietitians, and a simple portal for patients.",
        meta="Use this line in bios, ads, and the first sentence of long posts.",
    )
    post_box(
        doc,
        "Product facts (do not oversell)",
        "Standard is the full clinic workspace: charts, meal plans, food catalog, portal, messaging, calendar, tasks, invoices, analytics.\n\n"
        "Pro adds automations and AI assistance.\n\n"
        "14-day trial includes sample clients.\n\n"
        "Patients join with a clinic code.\n\n"
        "Invoices and quotations can be printed or saved as PDF. There is no built-in payment gateway.\n\n"
        "Food catalog includes USDA, Canadian Nutrient File, UK CoFID, and Lebanon 2021, plus practice custom foods and recipes.",
        meta="True claims only. If you are unsure, leave it out.",
    )

    add_heading_custom(doc, "Visual rules", 1)
    bullet(doc, "Mint, navy, white. Clean type. Illustrated clinic preview is fine.")
    bullet(doc, "No real patient names, no real dashboard, no real food-log screenshots.")
    bullet(doc, "No “AI dietitian,” no “patients pay in the app,” no competitor attacks.")
    bullet(doc, "Hashtags: 3 to 6. Suggested set: #Dietitian #NutritionClinic #PrivatePractice #PatientPortal #MealPlanning #DigitalHealth")

    # LinkedIn
    add_heading_custom(doc, "LinkedIn", 1)
    add_heading_custom(doc, "Launch post (full)", 2)
    post_box(
        doc,
        "Copy this into LinkedIn",
        "Most dietitians still run a clinic across spreadsheets, WhatsApp, and shared folders.\n\n"
        "The chart is in one place. The meal plan is in another. The patient logs food somewhere else. Invoices live in a third tool. Before a visit, you reconstruct the story by hand.\n\n"
        "That is not a clinic. It is a patchwork.\n\n"
        "We built Nutrition so the record, the plan, the logs, the messages, the visits, and the invoices sit on one client.\n\n"
        "For the dietitian\n"
        "• Client charts, clinical profile, measurements, and custom forms\n"
        "• Multi-week meal plans with automatic nutrition totals\n"
        "• Food catalog from USDA, Canada, UK, and Lebanon, plus your own foods and recipes\n"
        "• Calendar, tasks, secure messaging, and quotations or invoices you can print or save as PDF\n"
        "• Analytics for the practice\n\n"
        "For the patient\n"
        "• Join with a short clinic code\n"
        "• See the published plan, log food in one tap, track water, movement, sleep, habits, and weight\n"
        "• Message you and request or reschedule visits from a simple portal\n\n"
        "How it starts\n"
        "1. Share a clinic join code. The patient creates an account and appears on your roster.\n"
        "2. Build and publish the meal plan on the chart.\n"
        "3. Review logs before the next visit. Everything stays on the same record.\n\n"
        "Standard is the full clinic workspace. Pro adds automations and AI assistance.\n\n"
        "Start a 14-day trial with sample clients, then keep the plan that fits.\n\n"
        "[SITE]",
        meta="Best as a text post. First comment: hashtags. ~2,200 characters.",
        visual="Plain background or illustrated “clinic workspace + patient portal” graphic. No live UI.",
        cta="Start a 14-day trial",
        hashtags="#Dietitian #NutritionClinic #PrivatePractice #PatientPortal #DigitalHealth",
    )

    add_heading_custom(doc, "Shorter LinkedIn post", 2)
    post_box(
        doc,
        "Copy this into LinkedIn",
        "If your “clinic system” is Google Sheets + WhatsApp + a PDF meal plan, you do not have a system. You have three places to lose the same patient.\n\n"
        "We put the chart, the plan, the food log, messaging, visits, and invoices on one record. Patients join with a clinic code and follow care in a portal.\n\n"
        "14-day trial. Sample clients included.\n\n"
        "Dietitians: what is the one tool you still refuse to give up?\n\n"
        "[SITE]",
        meta="Use 2–3 days after the launch post. Ends with a question so comments start.",
        visual="Simple quote card: “Spreadsheets are not a clinic.”",
        cta="Start free trial",
        hashtags="#Dietitian #PrivatePractice #NutritionClinic",
    )

    add_heading_custom(doc, "LinkedIn comment to pin under launch", 2)
    post_box(
        doc,
        "First comment (pin it)",
        "How to try it\n"
        "1. Open [SITE]\n"
        "2. Start the 14-day trial (sample clients are already there)\n"
        "3. Share a clinic join code with a test patient, or use the samples\n\n"
        "Standard = full clinic. Pro = automations + AI.\n"
        "Questions welcome in the comments.",
        meta="Pin this so the main post stays clean.",
    )

    # X
    add_heading_custom(doc, "X (Twitter)", 1)
    add_heading_custom(doc, "Launch thread", 2)
    post_box(
        doc,
        "Tweet 1 of 5",
        "Spreadsheets are not a clinic.\n"
        "Chat threads are not a chart.\n"
        "A PDF meal plan is not follow-through.",
        meta="Hook. No link yet. Stay under 280 characters.",
        visual="None, or a 3-line graphic matching the text.",
    )
    post_box(
        doc,
        "Tweet 2 of 5",
        "Nutrition is a workspace for dietitians and a portal for patients.\n\n"
        "One client record. Plan, logs, messages, visits, invoices.",
        meta="Product definition.",
    )
    post_box(
        doc,
        "Tweet 3 of 5",
        "Patients join with a clinic code.\n"
        "You publish the meal plan.\n"
        "They log food, water, movement, sleep, and habits.\n"
        "You review before the visit.",
        meta="The three-step loop.",
    )
    post_box(
        doc,
        "Tweet 4 of 5",
        "Food data is not an afterthought. The catalog includes USDA, Canadian, UK, and Lebanon tables, plus your own foods and recipes. Nutrition totals calculate as you build.",
        meta="Proof, not hype.",
    )
    post_box(
        doc,
        "Tweet 5 of 5",
        "Standard = full clinic.\n"
        "Pro = automations + AI.\n"
        "14-day trial with sample clients.\n\n"
        "[SITE]",
        meta="Only tweet that carries the URL.",
        cta="Start trial",
        hashtags="#Dietitian #NutritionClinic",
    )

    add_heading_custom(doc, "Short posts to rotate", 2)
    shorts = [
        "Patients should not need three apps to follow a dietitian. One join code. One portal. One plan.",
        "Before the visit, open the chart. The logs are already there.",
        "Meal plans you can publish. Food logs you can actually review. That is the product.",
        "Automations and AI are on Pro. The clinic workspace is already on Standard.",
        "If you still email meal plans as attachments, try a 14-day trial. [SITE]",
        "A clinic code is enough. The patient creates an account and lands on your roster.",
        "Stop reconstructing care from WhatsApp the night before a visit.",
        "Standard is the full practice. Pro is for teams that want reminders and AI on top.",
    ]
    for i, text in enumerate(shorts, 1):
        post_box(
            doc,
            f"Short post {i}",
            text,
            meta="One per day. Add [SITE] only on posts 5, or on every third post.",
            visual="Plain text, or a single-sentence graphic.",
        )

    # Instagram
    add_heading_custom(doc, "Instagram", 1)
    add_heading_custom(doc, "Carousel post (10 slides)", 2)
    post_box(
        doc,
        "Caption (copy under the carousel)",
        "Most clinics are still a spreadsheet, a chat, and a folder. Nutrition keeps the chart, the meal plan, patient logs, messaging, visits, and invoices on one record. Patients join with a clinic code. 14-day trial. Link in bio.",
        meta="Keep the caption short. Details live on the slides.",
        visual="10 square slides, mint/navy, large type, no real UI.",
        cta="Link in bio",
        hashtags="#Dietitian #NutritionClinic #PrivatePractice #PatientPortal #MealPlanning",
    )

    add_heading_custom(doc, "Slide text (one idea per slide)", 2)
    slides = [
        ("01", "Spreadsheets are not a clinic."),
        ("02", "One workspace for the dietitian."),
        ("03", "A simple portal for the patient."),
        ("04", "Join with a clinic code."),
        ("05", "Publish the meal plan."),
        ("06", "Patients log food in one tap."),
        ("07", "Review progress before the visit."),
        ("08", "Catalog foods plus your own recipes."),
        ("09", "Standard or Pro. 14-day trial."),
        ("10", "Start at [SITE]"),
    ]
    add_table(doc, ["Slide", "On-slide copy"], [[a, b] for a, b in slides])

    add_heading_custom(doc, "Feed captions", 2)
    post_box(doc, "Caption A", "Add the client. Publish the plan. Review the logs.\n\nThree steps. One record.\n\nLink in bio.")
    post_box(doc, "Caption B", "For dietitians who are tired of reconstructing care from WhatsApp the night before a visit.\n\n14-day trial. Link in bio.")
    post_box(doc, "Caption C", "Patients follow the plan they were actually given. Not a screenshot from last month.\n\nLink in bio.")

    add_heading_custom(doc, "Stories (copy + sticker)", 2)
    stories = [
        ("1", "Poll", "Where do patient food logs live today?", "Options: WhatsApp / Spreadsheet / App / Nowhere"),
        ("2", "Educate", "Clinic join code", "Patients create their own account, then connect."),
        ("3", "Proof", "Today I published a 7-day plan", "Text only. No screenshot of a real client."),
        ("4", "Offer", "14-day trial is open", "Countdown sticker + link"),
        ("5", "Quiz", "What sits on one client record?", "Chart / Plan / Logs / All three (correct: all three)"),
        ("6", "Plans", "Standard = clinic. Pro = automations + AI.", "Link sticker"),
        ("7", "Catalog", "Food data from USDA, Canada, UK, and Lebanon", "Plus your own foods."),
        ("8", "Trust", "We do not replace your clinical judgment.", "We keep the record in one place."),
        ("9", "CTA", "Start the trial", "Link sticker to [SITE]"),
        ("10", "Portal", "Patient portal: plan, log, message, book.", None),
        ("11", "Contrast", "3 tools → 1 clinic", "Before / after words only"),
        ("12", "FAQ", "Do patients need an invite link?", "No. A clinic code is enough."),
        ("13", "Honest", "Invoices: print or save as PDF.", "No payment gateway baked in."),
        ("14", "Social proof", "Clinic quote", "Only with written permission."),
        ("15", "Reminder", "Trial includes sample clients", "You can click around today."),
    ]
    add_table(
        doc,
        ["#", "Type", "On-screen text", "Sticker / note"],
        [[a, b, c, d or "—"] for a, b, c, d in stories],
    )

    add_heading_custom(doc, "Reels (15–30 seconds)", 2)
    post_box(
        doc,
        "Reel 1 script",
        "HOOK on screen: Stop sending meal plans as PDFs.\n\n"
        "Voice (or captions): Patients join with a clinic code. You publish the plan. They log food. You review before the visit.\n\n"
        "End card: 14-day trial. Link in bio.",
        visual="Text-on-color cuts. No real dashboard.",
    )
    post_box(
        doc,
        "Reel 2 script",
        "HOOK: Your patient logged breakfast. Did you see it?\n\n"
        "Cut: It is on the chart.\n\n"
        "End: One record. Trial in bio.",
    )
    post_box(
        doc,
        "Reel 3 script",
        "Silent reel, big type:\n"
        "Problem: spreadsheet + chat + folder\n"
        "Join code → plan → log → visit\n"
        "CTA: Start the trial",
        meta="Works with music, no voiceover.",
    )

    # Facebook
    add_heading_custom(doc, "Facebook and dietitian groups", 1)
    post_box(
        doc,
        "Copy this (softer tone)",
        "Hi everyone. I built a clinic platform for dietitians because I was tired of seeing care split across sheets, chats, and folders.\n\n"
        "Patients join with a clinic code. You keep the chart and publish the meal plan. They log food and habits in a portal. Messaging, calendar, and invoices stay on the same record.\n\n"
        "Happy to share a 14-day trial if useful. Not here to spam. Questions welcome.\n\n"
        "[SITE]",
        meta="Groups: post as a person, not as an ad. Follow each group’s promo rules.",
        cta="Offer the trial only if someone asks, or once in the post.",
    )

    # Bios
    add_heading_custom(doc, "Profile bios", 1)
    post_box(doc, "LinkedIn headline", "Clinic workspace for dietitians. Patient portal included.")
    post_box(
        doc,
        "LinkedIn about (short)",
        "Nutrition is a clinic platform: charts, meal plans, food logs, messaging, visits, and invoices on one client record. Patients join with a clinic code. Standard or Pro. 14-day trial.",
    )
    post_box(
        doc,
        "Instagram bio",
        "Clinic workspace for dietitians\nPatient portal · join with a code\n14-day trial ↓",
        meta="Put [SITE] in the link sticker / bio URL.",
    )
    post_box(
        doc,
        "X bio",
        "Clinic workspace for dietitians. Portal for patients. Chart, plan, logs, visits. 14-day trial.",
    )

    # Calendar
    add_heading_custom(doc, "30-day posting calendar", 1)
    p = doc.add_paragraph()
    r = p.add_run("Post once on the main channel that day. Stories can run on Instagram the same day without a new feed post.")
    set_run(r, size=11, italic=True, color=MUTED)

    calendar = [
        ("1", "Launch", "LinkedIn long + IG carousel + X thread"),
        ("2", "Join code", "Story poll + short X"),
        ("3", "Meal plans", "Carousel: draft → publish → portal"),
        ("4", "Patient logging", "Reel: one-tap log"),
        ("5", "Food catalog", "What’s in the database (sources only)"),
        ("6", "Before the visit", "LinkedIn: open the chart, logs are there"),
        ("7", "Standard vs Pro", "Simple graphic"),
        ("8", "Messaging", "Care questions should not live in personal WhatsApp"),
        ("9", "Calendar", "Day / week / month"),
        ("10", "Invoices", "Print or save as PDF. No fake instant pay."),
        ("11", "Clinical profile", "Default questions every patient gets"),
        ("12", "Custom forms", "Separate from the default chart"),
        ("13", "Habits", "Assign, patient tracks"),
        ("14", "Trial reminder", "Sample clients included"),
        ("15", "Patient home", "Next visit, today’s log, plan snapshot"),
        ("16", "Measurements", "Weight, BMI, composition, labs over time"),
        ("17", "Automations (Pro)", "Reminders, inactivity, overdue invoices"),
        ("18", "AI (Pro)", "Assistance on the chart. Not a replacement."),
        ("19", "Analytics", "Practice pulse, not vanity metrics"),
        ("20", "Founder note", "Why we built it (your voice)"),
        ("21", "FAQ", "Join code, trial length, Standard vs Pro"),
        ("22", "Regional foods", "Catalog includes Lebanon 2021"),
        ("23", "Recipes", "Starter library + your own meals"),
        ("24", "Privacy-minded", "Professional record, not a public feed"),
        ("25", "Comparison", "Patchwork vs one record (no competitor names)"),
        ("26", "Patient story", "Anonymized, permission only"),
        ("27", "Clinic ops", "Tasks, due today, overdue"),
        ("28", "Social proof", "First real clinic quote"),
        ("29", "How to start", "Register → trial → add a real client"),
        ("30", "Recap + CTA", "Thread + LinkedIn + stories"),
    ]
    add_table(doc, ["Day", "Theme", "Format"], calendar)

    add_heading_custom(doc, "Weekly rhythm (after day 30)", 1)
    add_table(
        doc,
        ["Day", "Channel", "Job"],
        [
            ("Mon", "LinkedIn", "One professional insight or founder note"),
            ("Tue", "Instagram story", "Poll or FAQ"),
            ("Wed", "X", "Short product line, no thread"),
            ("Thu", "Instagram feed", "Carousel or reel"),
            ("Fri", "LinkedIn or groups", "Trial reminder or question"),
            ("Sat–Sun", "Stories only", "Light, or skip"),
        ],
    )

    add_heading_custom(doc, "Do not say", 1)
    bullet(doc, "“AI dietitian” or “replaces the clinician.”")
    bullet(doc, "“Patients pay in the app” or “automatic insurance billing.”")
    bullet(doc, "“Unlimited everything” unless the live plan really says so.")
    bullet(doc, "Real names, real labs, or a live clinic screenshot.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Replace [SITE] everywhere before you publish.")
    set_run(r, size=11, bold=True, color=TEAL)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
